"""Export a NewsletterRun as CSV (raw articles), JSON (full run), or DOCX (newsletter)."""
from __future__ import annotations

import csv
import io
import json

from docx import Document
from docx.shared import Pt

from backend.models.schemas import NewsletterRun


def to_csv(run: NewsletterRun) -> bytes:
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow([
        "title", "url", "source_domain", "credibility_tier", "published_date",
        "is_relevant", "relevance_score", "deal_type", "companies", "deal_amount",
        "one_line_summary", "also_reported_by",
    ])
    for article in run.raw_articles:
        writer.writerow([
            article.title,
            article.url,
            article.source_domain,
            article.credibility_tier.value if article.credibility_tier else "",
            article.published_date or "",
            article.is_relevant if article.is_relevant is not None else "",
            article.relevance_score if article.relevance_score is not None else "",
            article.deal_type or "",
            "; ".join(article.companies),
            article.deal_amount or "",
            article.one_line_summary or "",
            "; ".join(article.also_reported_by),
        ])
    return buffer.getvalue().encode("utf-8")


def to_json(run: NewsletterRun) -> bytes:
    return json.dumps(json.loads(run.model_dump_json()), indent=2).encode("utf-8")


def to_docx(run: NewsletterRun) -> bytes:
    document = Document()
    document.add_heading("FMCG Deal Intelligence Newsletter", level=0)
    if run.created_at:
        meta = document.add_paragraph(f"Generated: {run.created_at.isoformat()}")
        meta.runs[0].font.size = Pt(9)

    for section in run.newsletter_sections:
        document.add_heading(section.title, level=1)
        for deal in section.deals:
            document.add_heading(deal.headline, level=2)
            document.add_paragraph(deal.summary)
            details = document.add_paragraph()
            details.add_run("Companies: ").bold = True
            details.add_run(", ".join(deal.companies))
            if deal.deal_amount:
                amount = document.add_paragraph()
                amount.add_run("Deal size: ").bold = True
                amount.add_run(deal.deal_amount)
            sources = document.add_paragraph()
            sources.add_run("Sources: ").bold = True
            sources.add_run(f"{', '.join(deal.sources)} (credibility: {deal.credibility_tier.value})")

    if not run.newsletter_sections and run.newsletter_markdown:
        document.add_paragraph(run.newsletter_markdown)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


EXPORTERS = {"csv": to_csv, "json": to_json, "docx": to_docx}

MIME_TYPES = {
    "csv": "text/csv",
    "json": "application/json",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def export_run(run: NewsletterRun, fmt: str) -> bytes:
    if fmt not in EXPORTERS:
        raise ValueError(f"Unsupported export format: {fmt}")
    return EXPORTERS[fmt](run)
